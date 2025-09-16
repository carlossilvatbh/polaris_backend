"""
Document Processor - Processamento de documentos para RAG
Extração de texto de PDF, Word e outros formatos
"""

import logging
from typing import Dict, Any, List, Optional
from pathlib import Path

# Imports seguros com fallback
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

from .utils import RAGUtils

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Processador de documentos jurídicos para RAG"""
    
    def __init__(self):
        self.supported_formats = self._check_supported_formats()
        logger.info(f"DocumentProcessor inicializado. Formatos suportados: {self.supported_formats}")
    
    def _check_supported_formats(self) -> List[str]:
        """Verifica quais formatos estão disponíveis"""
        formats = ['.txt']  # Sempre disponível
        
        if PDF_AVAILABLE:
            formats.extend(['.pdf'])
        
        if DOCX_AVAILABLE:
            formats.extend(['.docx', '.doc'])
        
        return formats
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Processa um documento e extrai texto + metadados
        
        Args:
            file_path: Caminho para o arquivo
        
        Returns:
            Dict com text, metadata e chunks
        """
        try:
            # Valida arquivo
            is_valid, message = RAGUtils.validate_file_path(file_path)
            if not is_valid:
                return {
                    'success': False,
                    'error': message,
                    'text': '',
                    'metadata': {},
                    'chunks': []
                }
            
            path = Path(file_path)
            extension = path.suffix.lower()
            
            # Extrai texto baseado no formato
            if extension == '.pdf':
                text = self._extract_pdf_text(file_path)
            elif extension in ['.docx', '.doc']:
                text = self._extract_docx_text(file_path)
            elif extension == '.txt':
                text = self._extract_txt_text(file_path)
            else:
                return {
                    'success': False,
                    'error': f"Formato não suportado: {extension}",
                    'text': '',
                    'metadata': {},
                    'chunks': []
                }
            
            if not text.strip():
                return {
                    'success': False,
                    'error': "Nenhum texto extraído do documento",
                    'text': '',
                    'metadata': {},
                    'chunks': []
                }
            
            # Limpa e normaliza texto
            text = RAGUtils.clean_text(text)
            
            # Extrai metadados
            metadata = RAGUtils.extract_metadata_from_path(file_path)
            metadata.update({
                'text_length': len(text),
                'word_count': len(text.split()),
                'processed_by': 'DocumentProcessor',
                'extraction_method': self._get_extraction_method(extension)
            })
            
            # Gera chunks inteligentes
            chunks = RAGUtils.intelligent_chunk_text(text)
            
            # Adiciona metadados aos chunks
            for chunk in chunks:
                chunk['source_file'] = path.name
                chunk['source_path'] = file_path
                chunk['document_metadata'] = metadata
            
            logger.info(f"Documento processado: {path.name} - {len(chunks)} chunks gerados")
            
            return {
                'success': True,
                'text': text,
                'metadata': metadata,
                'chunks': chunks,
                'source_file': path.name
            }
            
        except Exception as e:
            error_msg = f"Erro ao processar documento {file_path}: {str(e)}"
            logger.error(error_msg)
            return {
                'success': False,
                'error': error_msg,
                'text': '',
                'metadata': {},
                'chunks': []
            }
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Alias para process_document. Extrai texto de um arquivo.
        Mantém compatibilidade com interface esperada pelos testes.
        
        Args:
            file_path: Caminho para o arquivo
            
        Returns:
            Dict com resultado da extração
        """
        return self.process_document(file_path)
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extrai texto de PDF usando múltiplas bibliotecas"""
        text = ""
        
        # Prioridade 1: pdfplumber (melhor para tabelas e layout)
        if PDF_AVAILABLE:
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n\n=== PÁGINA {page_num + 1} ===\n\n"
                            text += page_text
                
                if text.strip():
                    logger.debug(f"PDF extraído com pdfplumber: {len(text)} chars")
                    return text
            except Exception as e:
                logger.warning(f"Falha no pdfplumber: {str(e)}")
        
        # Prioridade 2: PyMuPDF (boa qualidade geral)
        if PYMUPDF_AVAILABLE:
            try:
                import fitz
                doc = fitz.open(file_path)
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    page_text = page.get_text()
                    if page_text:
                        text += f"\n\n=== PÁGINA {page_num + 1} ===\n\n"
                        text += page_text
                doc.close()
                
                if text.strip():
                    logger.debug(f"PDF extraído com PyMuPDF: {len(text)} chars")
                    return text
            except Exception as e:
                logger.warning(f"Falha no PyMuPDF: {str(e)}")
        
        # Prioridade 3: PyPDF2 (fallback)
        if PDF_AVAILABLE:
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n\n=== PÁGINA {page_num + 1} ===\n\n"
                            text += page_text
                
                if text.strip():
                    logger.debug(f"PDF extraído com PyPDF2: {len(text)} chars")
                    return text
            except Exception as e:
                logger.warning(f"Falha no PyPDF2: {str(e)}")
        
        raise Exception("Nenhuma biblioteca PDF disponível ou falha na extração")
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extrai texto de documentos Word"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx não está instalado")
        
        try:
            from docx import Document
            doc = Document(file_path)
            
            text_parts = []
            
            # Extrai texto dos parágrafos
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extrai texto das tabelas
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_texts.append(cell.text.strip())
                    if row_texts:
                        text_parts.append(" | ".join(row_texts))
            
            text = "\n\n".join(text_parts)
            logger.debug(f"DOCX extraído: {len(text)} chars")
            return text
            
        except Exception as e:
            raise Exception(f"Erro ao extrair texto do DOCX: {str(e)}")
    
    def _extract_txt_text(self, file_path: str) -> str:
        """Extrai texto de arquivos TXT"""
        try:
            # Tenta diferentes encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                        logger.debug(f"TXT extraído com encoding {encoding}: {len(text)} chars")
                        return text
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Não foi possível decodificar o arquivo de texto")
            
        except Exception as e:
            raise Exception(f"Erro ao extrair texto do TXT: {str(e)}")
    
    def _get_extraction_method(self, extension: str) -> str:
        """Retorna o método de extração usado"""
        if extension == '.pdf':
            if PDF_AVAILABLE:
                return "pdfplumber/PyMuPDF/PyPDF2"
            else:
                return "não_disponível"
        elif extension in ['.docx', '.doc']:
            return "python-docx" if DOCX_AVAILABLE else "não_disponível"
        elif extension == '.txt':
            return "built-in"
        else:
            return "não_suportado"
    
    def process_multiple_documents(self, file_paths: List[str]) -> Dict[str, Any]:
        """
        Processa múltiplos documentos em batch
        
        Args:
            file_paths: Lista de caminhos para arquivos
        
        Returns:
            Dict com resultados consolidados
        """
        results = {
            'successful': [],
            'failed': [],
            'total_chunks': 0,
            'total_text_length': 0,
            'processing_summary': {}
        }
        
        for file_path in file_paths:
            try:
                result = self.process_document(file_path)
                
                if result['success']:
                    results['successful'].append(result)
                    results['total_chunks'] += len(result['chunks'])
                    results['total_text_length'] += len(result['text'])
                else:
                    results['failed'].append({
                        'file_path': file_path,
                        'error': result['error']
                    })
                
            except Exception as e:
                results['failed'].append({
                    'file_path': file_path,
                    'error': str(e)
                })
        
        results['processing_summary'] = {
            'total_files': len(file_paths),
            'successful_files': len(results['successful']),
            'failed_files': len(results['failed']),
            'success_rate': len(results['successful']) / len(file_paths) if file_paths else 0,
            'average_chunks_per_file': (results['total_chunks'] / len(results['successful']) 
                                      if results['successful'] else 0),
            'total_text_length': results['total_text_length']
        }
        
        logger.info(f"Processamento em batch concluído: {results['processing_summary']}")
        return results
    
    def get_supported_formats(self) -> List[str]:
        """Retorna lista de formatos suportados"""
        return self.supported_formats.copy()
    
    def check_dependencies(self) -> Dict[str, bool]:
        """Verifica status das dependências"""
        return {
            'PyPDF2': PDF_AVAILABLE,
            'pdfplumber': PDF_AVAILABLE,
            'python-docx': DOCX_AVAILABLE,
            'PyMuPDF': PYMUPDF_AVAILABLE,
            'txt_support': True  # Sempre disponível
        }
